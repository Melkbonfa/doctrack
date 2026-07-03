# -*- coding: utf-8 -*-
"""Gera o relatorio de levantamento de custos de desenvolvimento do dashboard de TI,
em Word (.docx), no mesmo padrao visual dos relatorios da Loccus."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

ROXO = RGBColor(0x6A, 0x1B, 0x9A)
ROXO_HEX = "6A1B9A"
CINZA = RGBColor(0x44, 0x44, 0x44)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
VERDE = "2E7D32"
CINZA_CLARO = "EDE7F0"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = ROXO
    return p


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p


def body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i, part in enumerate(text.split("**")):
        r = p.add_run(part)
        r.font.size = Pt(11)
        if i % 2 == 1:
            r.bold = True
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    for i, part in enumerate(text.split("**")):
        r = p.add_run(part)
        r.font.size = Pt(11)
        if i % 2 == 1:
            r.bold = True
    return p


def brl(v):
    s = f"{v:,.2f}"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return "R$ " + s


# ---------------- DADOS ----------------
MESES = 5  # fev, mar, abr, mai, jun / 2026
itens = [
    # (categoria, descricao, mensal, meses, unico)
    ("Claude (Anthropic)", "Assinatura mensal da plataforma de IA", 118.0, MESES, 0.0),
    ("Claude (Anthropic)", "Creditos/tokens adicionais (compra avulsa)", 0.0, 0, 250.0),
    ("Antigravity", "Assinatura mensal da ferramenta de IA", 99.0, MESES, 0.0),
    ("Hospedagem", "Hospedagem da aplicacao (servidor web)", 20.0, MESES, 0.0),
    ("Banco de dados", "Servico de banco de dados PostgreSQL", 10.0, MESES, 0.0),
]


def total_item(mensal, meses, unico):
    return mensal * meses + unico


total_geral = sum(total_item(m, n, u) for _, _, m, n, u in itens)
mensal_recorrente = sum(m for _, _, m, n, u in itens if n)

# ---------------- CAPA ----------------
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Levantamento de Custos")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = ROXO
p = doc.add_paragraph()
r = p.add_run("Desenvolvimento do Dashboard de TI")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p = doc.add_paragraph()
r = p.add_run("Plataformas de IA, hospedagem e infraestrutura")
r.font.size = Pt(13)
r.font.color.rgb = CINZA

doc.add_paragraph()
meta = [
    ("Objeto", "Custos de ferramentas e infraestrutura do projeto"),
    ("Periodo", "Fevereiro a Junho de 2026 (5 meses)"),
    ("Moeda", "Real (R$)"),
    ("Custo total no periodo", brl(total_geral)),
    ("Custo recorrente mensal", brl(mensal_recorrente) + " / mes"),
    ("Data do levantamento", "22/06/2026"),
]
t = doc.add_table(rows=len(meta), cols=2)
t.style = "Table Grid"
for i, (k, v) in enumerate(meta):
    set_cell_text(t.cell(i, 0), k, bold=True, color=CINZA, size=10)
    shade(t.cell(i, 0), CINZA_CLARO)
    set_cell_text(t.cell(i, 1), v, size=10)
    t.cell(i, 0).width = Cm(5.0)
    t.cell(i, 1).width = Cm(11.0)

doc.add_page_break()

# ---------------- 1. RESUMO ----------------
h1("1. Resumo Executivo")
body(
    f"Este documento consolida os custos de desenvolvimento do **Dashboard de TI**, considerando as "
    f"plataformas de inteligencia artificial utilizadas, a hospedagem da aplicacao e o servico de banco de "
    f"dados. O periodo abrange **fevereiro a junho de 2026 (5 meses)**, durante o qual as assinaturas foram "
    f"mantidas de forma continua."
)
body(
    f"O custo **total acumulado no periodo** foi de **{brl(total_geral)}**, dos quais "
    f"**{brl(mensal_recorrente)} por mes** correspondem a despesas recorrentes (assinaturas e infraestrutura) e "
    f"**{brl(250.0)}** a uma compra avulsa de creditos adicionais do Claude."
)

# ---------------- 2. DETALHAMENTO ----------------
h1("2. Detalhamento por Item")
cab = ("Categoria", "Descricao", "Mensal", "Meses", "Total")
linhas = [cab]
for cat, desc, m, n, u in itens:
    mensal_txt = brl(m) if m else "-"
    meses_txt = str(n) if n else "unico"
    linhas.append((cat, desc, mensal_txt, meses_txt, brl(total_item(m, n, u))))

tb = doc.add_table(rows=len(linhas) + 1, cols=5)
tb.style = "Table Grid"
larguras = [Cm(3.4), Cm(6.2), Cm(2.4), Cm(1.6), Cm(2.6)]
for i, linha in enumerate(linhas):
    for j, val in enumerate(linha):
        if i == 0:
            set_cell_text(tb.cell(i, j), val, bold=True, color=BRANCO, size=10,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            shade(tb.cell(i, j), ROXO_HEX)
        else:
            al = WD_ALIGN_PARAGRAPH.CENTER if j in (2, 3, 4) else None
            set_cell_text(tb.cell(i, j), val, size=10, align=al)
        tb.cell(i, j).width = larguras[j]
# linha de total
ult = len(linhas)
set_cell_text(tb.cell(ult, 0), "TOTAL GERAL", bold=True, color=BRANCO, size=10)
shade(tb.cell(ult, 0), VERDE)
for j in (1, 2, 3):
    set_cell_text(tb.cell(ult, j), "", size=10)
    shade(tb.cell(ult, j), VERDE)
set_cell_text(tb.cell(ult, 4), brl(total_geral), bold=True, color=BRANCO, size=11,
              align=WD_ALIGN_PARAGRAPH.CENTER)
shade(tb.cell(ult, 4), VERDE)
for j in range(5):
    tb.cell(ult, j).width = larguras[j]

doc.add_paragraph()

# ---------------- 3. POR CATEGORIA ----------------
h1("3. Totais por Categoria")
porcat = {}
for cat, desc, m, n, u in itens:
    porcat[cat] = porcat.get(cat, 0.0) + total_item(m, n, u)

linhas2 = [("Categoria", "Total no periodo", "% do total")]
for cat, val in sorted(porcat.items(), key=lambda x: -x[1]):
    pct = (val / total_geral * 100) if total_geral else 0
    linhas2.append((cat, brl(val), f"{pct:.1f}%"))

tc = doc.add_table(rows=len(linhas2), cols=3)
tc.style = "Table Grid"
for i, linha in enumerate(linhas2):
    for j, val in enumerate(linha):
        if i == 0:
            set_cell_text(tc.cell(i, j), val, bold=True, color=BRANCO, size=10,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            shade(tc.cell(i, j), ROXO_HEX)
        else:
            al = WD_ALIGN_PARAGRAPH.CENTER if j in (1, 2) else None
            set_cell_text(tc.cell(i, j), val, size=10, align=al)
    tc.cell(i, 0).width = Cm(6.0)
    tc.cell(i, 1).width = Cm(5.0)
    tc.cell(i, 2).width = Cm(5.0)

doc.add_paragraph()

# grafico de pizza do custo real por categoria
_cats = sorted(porcat.items(), key=lambda x: -x[1])
fig, ax = plt.subplots(figsize=(5.2, 3.6))
cores_pie = ["#6A1B9A", "#9C4DCC", "#BA68C8", "#CE93D8"]
ax.pie([v for _, v in _cats], labels=[c for c, _ in _cats],
       autopct=lambda p: brl(p / 100 * total_geral).replace("R$ ", "R$\n"),
       colors=cores_pie[:len(_cats)], startangle=90,
       textprops={"fontsize": 8})
ax.set_title("Custo real por categoria", fontsize=11, color="#6A1B9A", weight="bold")
fig.tight_layout()
fig.savefig("_graf_pizza.png", dpi=150)
plt.close(fig)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture("_graf_pizza.png", width=Cm(11))

# ---------------- 4. MAO DE OBRA ESTIMADA ----------------
SEMANAS = 20  # fev a jun/2026
H_MIN, H_MAX = 400, 500          # 20-25h/sem x ~20 semanas
RH_MIN, RH_MAX = 50.0, 120.0     # faixa de mercado R$/h
H_MED, RH_MED = 450, 85.0

mo_min = H_MIN * RH_MIN
mo_max = H_MAX * RH_MAX
mo_med = H_MED * RH_MED

h1("4. Mao de Obra Estimada (Investimento de Esforco)")
body(
    "Este bloco e uma **estimativa de esforco** e nao representa desembolso financeiro (caixa). Serve para "
    "dimensionar o investimento total e o valor entregue. O desenvolvimento consumiu cerca de **20 a 25 horas "
    "por semana ao longo de ~20 semanas** (fev-jun/2026), divididas **metade em expediente convencional e "
    "metade fora do expediente**, totalizando aproximadamente **400 a 500 horas**."
)
body(
    "A valoracao usa a **faixa de mercado de R$ 50 a R$ 120 por hora** para desenvolvimento, apresentada em tres "
    "cenarios (minimo, medio e maximo)."
)

h2("Divisao das horas (cenario medio ~450h)")
linhas_h = [
    ("Contexto", "Horas (aprox.)", "Participacao"),
    ("Expediente convencional", "~225 h", "50%"),
    ("Fora do expediente", "~225 h", "50%"),
    ("Total", "~450 h", "100%"),
]
th = doc.add_table(rows=len(linhas_h), cols=3)
th.style = "Table Grid"
for i, linha in enumerate(linhas_h):
    for j, val in enumerate(linha):
        if i == 0:
            set_cell_text(th.cell(i, j), val, bold=True, color=BRANCO, size=10,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            shade(th.cell(i, j), ROXO_HEX)
        else:
            bold = (linha[0] == "Total")
            al = WD_ALIGN_PARAGRAPH.CENTER if j in (1, 2) else None
            set_cell_text(th.cell(i, j), val, bold=bold, size=10, align=al)
    th.cell(i, 0).width = Cm(6.0)
    th.cell(i, 1).width = Cm(5.0)
    th.cell(i, 2).width = Cm(5.0)

doc.add_paragraph()
h2("Cenarios de valoracao")
linhas_c = [
    ("Cenario", "Horas", "Valor-hora", "Mao de obra estimada"),
    ("Minimo", f"{H_MIN} h", brl(RH_MIN), brl(mo_min)),
    ("Medio", f"{H_MED} h", brl(RH_MED), brl(mo_med)),
    ("Maximo", f"{H_MAX} h", brl(RH_MAX), brl(mo_max)),
]
tcen = doc.add_table(rows=len(linhas_c), cols=4)
tcen.style = "Table Grid"
for i, linha in enumerate(linhas_c):
    for j, val in enumerate(linha):
        if i == 0:
            set_cell_text(tcen.cell(i, j), val, bold=True, color=BRANCO, size=10,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            shade(tcen.cell(i, j), ROXO_HEX)
        else:
            bold = (linha[0] == "Medio")
            al = WD_ALIGN_PARAGRAPH.CENTER if j in (1, 2, 3) else None
            set_cell_text(tcen.cell(i, j), val, bold=bold, size=10, align=al)
            if linha[0] == "Medio":
                shade(tcen.cell(i, j), CINZA_CLARO)
    tcen.cell(i, 0).width = Cm(3.5)
    tcen.cell(i, 1).width = Cm(3.0)
    tcen.cell(i, 2).width = Cm(3.5)
    tcen.cell(i, 3).width = Cm(6.0)

doc.add_paragraph()

# grafico de barras dos cenarios de mao de obra
fig, ax = plt.subplots(figsize=(5.6, 3.4))
labels_b = ["Minimo", "Medio", "Maximo"]
valores_b = [mo_min, mo_med, mo_max]
barras = ax.bar(labels_b, valores_b, color=["#BA68C8", "#6A1B9A", "#4A148C"])
ax.set_ylabel("R$", fontsize=9)
ax.set_title("Mao de obra estimada por cenario", fontsize=11, color="#6A1B9A", weight="bold")
for b, v in zip(barras, valores_b):
    ax.text(b.get_x() + b.get_width() / 2, v, brl(v), ha="center", va="bottom", fontsize=8)
ax.margins(y=0.18)
ax.spines[["top", "right"]].set_visible(False)
fig.tight_layout()
fig.savefig("_graf_barras.png", dpi=150)
plt.close(fig)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture("_graf_barras.png", width=Cm(12))
h2("Consolidado: custo real x investimento estimado")
linhas_t = [
    ("Bloco", "Natureza", "Valor"),
    ("A - Custo real (caixa)", "Desembolso efetivo (IA + infra)", brl(total_geral)),
    ("B - Mao de obra estimada", "Esforco valorado (cenario medio)", brl(mo_med)),
    ("Investimento total (A + B)", "Caixa + esforco (cenario medio)", brl(total_geral + mo_med)),
]
tt = doc.add_table(rows=len(linhas_t), cols=3)
tt.style = "Table Grid"
for i, linha in enumerate(linhas_t):
    for j, val in enumerate(linha):
        if i == 0:
            set_cell_text(tt.cell(i, j), val, bold=True, color=BRANCO, size=10,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            shade(tt.cell(i, j), ROXO_HEX)
        elif linha[0].startswith("Investimento"):
            set_cell_text(tt.cell(i, j), val, bold=True, color=BRANCO, size=10)
            shade(tt.cell(i, j), VERDE)
        else:
            set_cell_text(tt.cell(i, j), val, size=10)
    tt.cell(i, 0).width = Cm(5.5)
    tt.cell(i, 1).width = Cm(6.5)
    tt.cell(i, 2).width = Cm(4.0)

body(
    f"Em outras palavras: o desembolso real foi de **{brl(total_geral)}**, enquanto o **esforco de "
    f"desenvolvimento equivale a algo entre {brl(mo_min)} e {brl(mo_max)}** se fosse contratado no mercado "
    f"(referencia media de **{brl(mo_med)}**). Esse valor representa a economia gerada ao desenvolver o "
    f"dashboard internamente."
)

doc.add_paragraph()

# ---------------- 5. OBSERVACOES ----------------
h1("5. Observacoes")
bullet("Os valores de assinatura sao recorrentes e continuarao incidindo enquanto as plataformas forem mantidas.")
bullet(f"O custo recorrente atual e de **{brl(mensal_recorrente)} por mes** "
       f"(Claude, Antigravity, hospedagem e banco de dados).")
bullet("Os creditos adicionais do Claude (R$ 250) foram uma compra pontual para ampliar o limite de uso "
       "em um periodo de maior demanda.")
bullet("Projecao: mantido o ritmo atual, o custo recorrente equivale a "
       f"**{brl(mensal_recorrente * 12)} por ano**.")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Relatorio gerado em 22/06/2026. Documento de uso interno - editavel.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = CINZA

doc.save("Levantamento_Custos_Dashboard_TI.docx")
print("DOCX gerado com sucesso. Total:", brl(total_geral))

# ---------------- PDF ----------------
try:
    from docx2pdf import convert
    convert("Levantamento_Custos_Dashboard_TI.docx", "Levantamento_Custos_Dashboard_TI.pdf")
    print("PDF gerado com sucesso.")
except Exception as e:
    print("Falha ao gerar PDF:", e)
