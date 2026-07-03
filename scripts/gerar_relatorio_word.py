# -*- coding: utf-8 -*-
"""Gera o relatorio tecnico de avaliacao do site staging.loccus.com.br em Word (.docx),
com espacos reservados/legendas para insercao das imagens (prints)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROXO = RGBColor(0x6A, 0x1B, 0x9A)
ROXO_HEX = "6A1B9A"
CINZA = RGBColor(0x44, 0x44, 0x44)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
VERMELHO = "C62828"
LARANJA = "EF6C00"
AMARELO = "F9A825"
AZUL = "1565C0"
CINZA_CLARO = "EDE7F0"

doc = Document()

# Estilo base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def set_cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = ROXO
    p.space_before = Pt(12)
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    runs = []
    # suporte simples a **negrito**
    parts = text.split("**")
    for i, part in enumerate(parts):
        r = p.add_run(part)
        r.font.size = Pt(11)
        if i % 2 == 1:
            r.bold = True
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    parts = text.split("**")
    for i, part in enumerate(parts):
        r = p.add_run(part)
        r.font.size = Pt(11)
        if i % 2 == 1:
            r.bold = True
    return p

def placeholder_imagem(legenda):
    """Insere uma caixa cinza com instrucao para encaixar o print + legenda."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, "F2F2F2")
    cell.width = Cm(15)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n[  INSERIR IMAGEM AQUI  ]\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(legenda)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = CINZA
    # borda
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run("Figura — " + legenda)
    rc.italic = True
    rc.font.size = Pt(9)
    rc.font.color.rgb = CINZA

# ---------------- CAPA ----------------
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Relatório Técnico de Avaliação")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = ROXO
p = doc.add_paragraph()
r = p.add_run("Inconsistências de Design e Funcionamento")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p = doc.add_paragraph()
r = p.add_run("Ambiente avaliado: staging.loccus.com.br")
r.font.size = Pt(13)
r.font.color.rgb = CINZA

doc.add_paragraph()
meta = [
    ("Objeto", "Site institucional Loccus (ambiente de homologação/staging)"),
    ("Tipo de avaliação", "Inspeção de design (UI/UX) e funcionamento (QA funcional)"),
    ("Métodos", "Análise de HTML/conteúdo, navegação assistida (desktop), console e validação de formulário, complementada por revisão de UX da equipe"),
    ("Data da avaliação", "17/06/2026"),
    ("Páginas avaliadas", "Home (PT/EN), Listagem de Produtos, Blog, Contato"),
]
t = doc.add_table(rows=len(meta), cols=2)
t.style = "Table Grid"
for i, (k, v) in enumerate(meta):
    set_cell_text(t.cell(i, 0), k, bold=True, color=CINZA, size=10)
    shade(t.cell(i, 0), CINZA_CLARO)
    set_cell_text(t.cell(i, 1), v, size=10)
    t.cell(i, 0).width = Cm(4.5)
    t.cell(i, 1).width = Cm(11.5)

doc.add_page_break()

# ---------------- 1. RESUMO EXECUTIVO ----------------
h1("1. Resumo Executivo")
body("Esta avaliação verificou o site em ambiente de homologação (**staging.loccus.com.br**) quanto a "
     "inconsistências de design e de funcionamento. Foram combinadas a análise direta do HTML/conteúdo de "
     "múltiplas páginas, a navegação assistida no navegador (captura de tela, leitura do console JavaScript e "
     "teste de validação de formulário) e a **revisão de UX conduzida pela equipe**.")
body("No total foram consolidados **22 apontamentos**, organizados em dois blocos: **12 achados técnicos** "
     "(conteúdo, internacionalização, funcionamento e SEO) e **10 apontamentos de experiência do usuário (UX)** "
     "levantados pela equipe. Os pontos mais críticos são de natureza editorial e de internacionalização: "
     "**texto de preenchimento (“Lorem ipsum”) publicado**, **banner de site de desenvolvimento do WPML**, "
     "**variáveis de template não renderizadas** no aviso de cookies e **tradução incompleta da versão em inglês**. "
     "Em UX, destacam-se **legibilidade prejudicada por efeito blur**, **proporção inconsistente das imagens de "
     "produto**, **botões sem mudança de estado** e **pop-up exibido cedo demais**.")
body("Recomenda-se tratar os itens de severidade **Alta** antes da promoção do ambiente para produção, por serem "
     "visíveis ao público e afetarem a credibilidade institucional e o SEO.")

h2("Síntese por severidade")
sev = [
    ("Severidade", "Qtd.", "Itens"),
    ("ALTA", "3", "Lorem ipsum publicado; banner WPML de desenvolvimento; variáveis de template no banner de cookies."),
    ("MÉDIA", "5", "Tradução EN incompleta; bug de layout no formulário; exceções de JS; rota /en/ 404; ausência de hreflang."),
    ("BAIXA", "4", "Inconsistência tipográfica; cards de Aplicações desalinhados; datas do blog; imagens sem alt."),
    ("UX (equipe)", "10", "Ícones redundantes; estado do botão do esquema; botão VEJA MAIS desproporcional; blur (2x); caixa vazia; carrossel sem setas; pop-up precoce; proporção/desalinhamento de produtos."),
]
ts = doc.add_table(rows=len(sev), cols=3)
ts.style = "Table Grid"
cores = {"ALTA": VERMELHO, "MÉDIA": LARANJA, "BAIXA": AMARELO, "UX (equipe)": AZUL}
for i, (a, b, c) in enumerate(sev):
    if i == 0:
        for j, txt in enumerate((a, b, c)):
            set_cell_text(ts.cell(i, j), txt, bold=True, color=BRANCO, size=10)
            shade(ts.cell(i, j), ROXO_HEX)
    else:
        set_cell_text(ts.cell(i, 0), a, bold=True, color=BRANCO, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(ts.cell(i, 0), cores[a])
        set_cell_text(ts.cell(i, 1), b, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(ts.cell(i, 2), c, size=10)
    ts.cell(i, 0).width = Cm(3)
    ts.cell(i, 1).width = Cm(1.5)
    ts.cell(i, 2).width = Cm(11.5)

# ---------------- 2. METODOLOGIA ----------------
h1("2. Escopo e Metodologia")
body("A avaliação foi conduzida sobre o ambiente de homologação e não incluiu testes de carga, segurança ou "
     "compatibilidade entre múltiplos navegadores. As técnicas empregadas foram:")
for it in [
    "**Análise de conteúdo/HTML:** leitura do código e do conteúdo renderizado das páginas Home, Blog, Produtos e Contato (PT e EN).",
    "**Navegação assistida (desktop):** percurso visual da página inicial, seções, listagem de produtos e rodapé.",
    "**Console JavaScript:** leitura das mensagens de console durante o carregamento da Home.",
    "**Teste funcional de formulário:** submissão do formulário de contato em branco para verificar a validação.",
    "**Internacionalização:** comparação da versão PT com a versão EN (?lang=en).",
    "**Revisão de UX da equipe:** apontamentos qualitativos de experiência do usuário registrados com capturas de tela.",
]:
    bullet(it)
body("**Limitação:** a validação de layout responsivo (mobile) não pôde ser confirmada de forma confiável nesta "
     "avaliação. Recomenda-se reteste específico de mobile.")

doc.add_page_break()

# ---------------- 3. ACHADOS TECNICOS ----------------
h1("3. Achados Técnicos")

def achado(cod, titulo, sevtxt, sevcor, categoria, descricao, evidencia, recomendacao, legenda_img=None):
    h2(f"{cod}  {titulo}")
    # linha severidade/categoria
    tb = doc.add_table(rows=1, cols=2)
    set_cell_text(tb.cell(0, 0), sevtxt, bold=True, color=BRANCO, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(tb.cell(0, 0), sevcor)
    set_cell_text(tb.cell(0, 1), "Categoria: " + categoria, size=10)
    tb.cell(0, 0).width = Cm(3)
    tb.cell(0, 1).width = Cm(13)
    # bloco
    info = [("Descrição", descricao), ("Evidência", evidencia), ("Recomendação", recomendacao)]
    bt = doc.add_table(rows=len(info), cols=2)
    bt.style = "Table Grid"
    for i, (k, v) in enumerate(info):
        set_cell_text(bt.cell(i, 0), k, bold=True, color=CINZA, size=9)
        shade(bt.cell(i, 0), CINZA_CLARO)
        set_cell_text(bt.cell(i, 1), v, size=10)
        bt.cell(i, 0).width = Cm(3)
        bt.cell(i, 1).width = Cm(13)
    if legenda_img:
        placeholder_imagem(legenda_img)
    doc.add_paragraph()

tecnicos = [
    ("3.1", "Texto “Lorem ipsum” publicado na Home e no Blog", "ALTA", VERMELHO, "Conteúdo / Editorial",
     "A seção “Fique por dentro” da Home e o Blog exibem um post real com título “Lorem ipsum dolor” e corpo de texto de preenchimento, além do post padrão “Hello world!” do WordPress, que nunca foi removido.",
     "Confirmado visualmente. Texto: “Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nulla fermentum arcu metus...”.",
     "Remover/despublicar os posts de exemplo e substituir por conteúdo real antes da publicação.",
     "Seção “Fique por dentro” com post “Lorem ipsum dolor”."),
    ("3.2", "Banner de “site de desenvolvimento” do WPML", "ALTA", VERMELHO, "Configuração / SEO",
     "O conteúdo apresenta o aviso do WPML indicando que o site está registrado como ambiente de desenvolvimento.",
     "Texto presente no HTML: “Este site está registrado em wpml.org como um site de desenvolvimento.”",
     "Aplicar a chave de produção do WPML para remover o banner antes de promover a produção.",
     None),
    ("3.3", "Variáveis de template não renderizadas no banner de cookies", "ALTA", VERMELHO, "Defeito de renderização",
     "O banner de gestão de cookies exibe marcadores de template em vez dos valores reais.",
     "Ocorrências de {title} (3x) e {vendor_count} (ex.: “Gerenciar {vendor_count} fornecedores”).",
     "Revisar a configuração/integração do plugin de consentimento (Complianz) para que as variáveis sejam preenchidas.",
     None),
    ("3.4", "Tradução incompleta da versão em inglês", "MÉDIA", LARANJA, "Internacionalização (i18n)",
     "Na versão em inglês (?lang=en), diversos blocos permanecem em português.",
     "Confirmado: bloco Newsletter (“Newsletter” / “Cadastre-se e receba nossas informações” / “Inscreva-se”), cabeçalho “Contato” no rodapé, “Loccus® - Todos os direitos reservados.” e botão “LEIA MAIS” do blog.",
     "Completar as strings de tradução (WPML), incluindo widgets de rodapé, newsletter e rótulos de botões.",
     "Versão EN com bloco Newsletter ainda em português."),
    ("3.5", "Defeito de layout no formulário de contato ao validar", "MÉDIA", LARANJA, "Funcional / UI",
     "Ao submeter o formulário “Fale Conosco” com erro de validação, o botão ENVIAR é reposicionado para cima e fica sobreposto à seção “Intenção”. A validação destacou apenas o telefone, sem marcar os campos obrigatórios Nome e E-mail.",
     "Confirmado em teste de submissão em branco: “Digite um número de telefone.” e “Um ou mais campos possuem um erro...”, com o botão sobreposto.",
     "Corrigir o CSS das mensagens de erro (evitar reflow que sobrepõe elementos) e validar todos os campos obrigatórios.",
     "Formulário com botão ENVIAR sobreposto após erro."),
    ("3.6", "Exceções de JavaScript no carregamento da Home", "MÉDIA", LARANJA, "Funcional / JS",
     "Foram registradas duas exceções de JavaScript no console durante o carregamento da página inicial.",
     "Mensagens do tipo EXCEPTION capturadas no console em staging.loccus.com.br.",
     "Investigar a origem das exceções (stack trace) e corrigir, pois podem afetar interações futuras.",
     None),
    ("3.7", "Rota /en/ retorna 404", "MÉDIA", LARANJA, "Internacionalização / Navegação",
     "O seletor de idioma utiliza query string (?lang=en, ?lang=es). A URL “limpa” /en/ retorna erro 404.",
     "Requisição a /en/ resultou em HTTP 404 Not Found.",
     "Garantir consistência de URLs de idioma e/ou redirecionar /en/ para o formato suportado.",
     None),
    ("3.8", "Ausência de tags hreflang", "MÉDIA", LARANJA, "SEO multilíngue",
     "O site oferece PT, EN e ES, mas não declara tags hreflang no cabeçalho do documento.",
     "Nenhuma tag <link rel=\"hreflang\"> identificada no <head>.",
     "Adicionar hreflang para cada idioma/variante, melhorando a indexação por mecanismos de busca.",
     None),
    ("3.9", "Inconsistência tipográfica entre seções", "BAIXA", AMARELO, "Design / Identidade visual",
     "Coexistem ao menos três estilos de título: o hero em caixa-alta pesada, títulos de seção em fonte geométrica (“Matrizes”, “Equipamentos”) e títulos como “Fluxo de Trabalho” / “Sobre nós” em fonte fina e arredondada.",
     "Confirmado visualmente em diferentes seções da Home.",
     "Padronizar a hierarquia tipográfica conforme um guia de estilo único (famílias, pesos e usos).",
     None),
    ("3.10", "Cards de “Aplicações” desalinhados e com baixa legibilidade", "BAIXA", AMARELO, "Design / Acessibilidade",
     "As imagens da seção Aplicações ficam escalonadas em alturas diferentes e alguns rótulos brancos sobre fundo claro têm contraste insuficiente, com texto colado à borda inferior.",
     "Confirmado visualmente (ex.: “Precision Agriculture” e “Forensics and Human Identification”).",
     "Padronizar alturas/alinhamento dos cards e reforçar o contraste do rótulo (overlay/sombra) para WCAG AA.",
     None),
    ("3.11", "Datas do blog em formato incomum e data futura", "BAIXA", AMARELO, "Conteúdo",
     "As datas do blog usam o formato “02 / 03 / 2026” (com espaços) e há post com data associada a conteúdo de exemplo.",
     "Confirmado visualmente no card de blog.",
     "Padronizar o formato de data e revisar datas/conteúdos de exemplo.",
     None),
    ("3.12", "Imagens sem atributo alt", "BAIXA", AMARELO, "Acessibilidade / SEO",
     "Imagens do carrossel de produtos e do hero não apresentam texto alternativo.",
     "Identificado na análise de HTML das imagens (carrossel e hero).",
     "Adicionar textos alternativos descritivos a todas as imagens informativas.",
     None),
]
for a in tecnicos:
    achado(*a)

doc.add_page_break()

# ---------------- 4. APONTAMENTOS DE UX (EQUIPE) ----------------
h1("4. Apontamentos de UX (Revisão da Equipe)")
body("Esta seção consolida os apontamentos de experiência do usuário levantados pela equipe durante a navegação, "
     "com a respectiva captura de tela (a inserir) e a recomendação sugerida.")

def ux(cod, titulo, categoria, descricao, recomendacao, legenda_img):
    h2(f"{cod}  {titulo}")
    bt = doc.add_table(rows=3, cols=2)
    bt.style = "Table Grid"
    info = [("Categoria", categoria), ("Observação", descricao), ("Recomendação", recomendacao)]
    for i, (k, v) in enumerate(info):
        set_cell_text(bt.cell(i, 0), k, bold=True, color=CINZA, size=9)
        shade(bt.cell(i, 0), CINZA_CLARO)
        set_cell_text(bt.cell(i, 1), v, size=10)
        bt.cell(i, 0).width = Cm(3)
        bt.cell(i, 1).width = Cm(13)
    placeholder_imagem(legenda_img)
    doc.add_paragraph()

ux_itens = [
    ("4.1", "Dois ícones com a mesma função no header", "UX / Redundância",
     "Os ícones de “pessoa” e de “documento” no cabeçalho parecem apontar para a mesma função, gerando redundância e confusão.",
     "Unificar em um único ícone ou diferenciar claramente as funções (ex.: “Minha conta” vs. “Meu orçamento”) com rótulos/tooltips.",
     "Header com os dois ícones (pessoa + documento)."),
    ("4.2", "Esquema “Fluxo de Trabalho” – botão sem mudança de estado", "UI / Estado de componente",
     "O esquema interativo (Preparação/Extração/etc.) não foi bem avaliado; em especial, o botão/etapa ativa não muda de estado, apenas exibe uma linha embaixo, dificultando a percepção de seleção.",
     "Revisar o componente; aplicar estado ativo claro (cor de fundo, contraste, destaque) em vez de apenas sublinhado. Avaliar redesenho do esquema.",
     "Esquema “Fluxo de Trabalho” com etapa “Preparação” ativa."),
    ("4.3", "Botão “VEJA MAIS” desproporcional", "UI / Proporção",
     "O botão “VEJA MAIS” aparece grande demais e desconexo do contexto dos cards ao redor.",
     "Ajustar o tamanho/peso do botão para harmonizar com os cards e a hierarquia visual da seção.",
     "Seção de equipamentos com o botão “VEJA MAIS” grande."),
    ("4.4", "Texto pouco legível sobre efeito blur (Sobre nós)", "Legibilidade / Acessibilidade",
     "O card “Uma indústria focada em tornar a ciência mais acessível...” fica pouco legível devido ao efeito blur ao fundo.",
     "Aumentar o contraste do texto, reduzir a intensidade do blur ou aplicar fundo sólido/semitransparente atrás do texto.",
     "Card “Sobre nós” com texto sobre fundo desfocado."),
    ("4.5", "Caixa de texto vazia abaixo de “Fique por dentro”", "Layout / Conteúdo",
     "Há uma caixa de texto aparentemente vazia logo abaixo da seção “Fique por dentro”.",
     "Remover a caixa vazia ou preencher com o conteúdo previsto.",
     "Caixa vazia abaixo de “Fique por dentro”."),
    ("4.6", "Blur prejudica a leitura do título “Aplicações”", "Legibilidade",
     "O efeito blur aplicado ao fundo do título “Aplicações” compromete a leitura.",
     "Reduzir o blur e/ou aumentar o contraste entre título e fundo.",
     "Título “Aplicações” com efeito blur."),
    ("4.7", "Carrossel exige arrastar com o mouse", "UX / Navegação",
     "O carrossel de cards (“Simplificamos processos...”) exige arraste com o mouse; faltam controles (setas/botões) para girar os cards.",
     "Adicionar botões de navegação (setas anterior/próximo) e indicadores clicáveis ao carrossel.",
     "Carrossel “Simplificamos processos para que você foque...”."),
    ("4.8", "Pop-up “Saiba mais” exibido cedo demais", "UX / Timing de pop-up",
     "Ao clicar no Extracta Assist, abre-se o Extracta 16 já exibindo o pop-up “Saiba mais”, antes de o usuário ter tempo de conhecer o produto.",
     "Atrasar/condicionar a exibição do pop-up (ex.: após rolagem ou intenção de saída), evitando interrupção imediata.",
     "Pop-up “Saiba mais” sobre a página do Extracta 16."),
    ("4.9", "Proporção das imagens de produto inconsistente", "UI / Proporção de imagem",
     "Na listagem de produtos, equipamentos grandes aparecem pequenos e equipamentos pequenos aparecem grandes, sem escala/proporção padronizada entre os cards.",
     "Definir proporção/escala consistente para as fotos de produto (mesma área visual e enquadramento padronizado).",
     "Listagem de produtos com escalas de imagem divergentes."),
    ("4.10", "Desalinhamentos nos cards de produto", "UI / Alinhamento",
     "Os cards de produto apresentam alturas e posições variando entre si, gerando desalinhamento visual.",
     "Padronizar a altura dos cards e o alinhamento de imagem, título, descrição e botões.",
     "Grade de produtos com cards desalinhados."),
]
for u in ux_itens:
    ux(*u)

doc.add_page_break()

# ---------------- 5. CONCLUSAO ----------------
h1("5. Conclusão e Plano de Ação")
body("O ambiente avaliado encontra-se funcional em sua estrutura principal, com navegação, seletor de idioma "
     "(via query string) e formulário de contato operantes. Contudo, há pendências editoriais, de "
     "internacionalização e de experiência do usuário que devem ser resolvidas antes da promoção a produção, "
     "em especial os itens de severidade Alta, por serem visíveis ao público.")
h2("Plano de ação sugerido (por prioridade)")
for it in [
    "**1. Imediato (severidade Alta):** remover conteúdo “Lorem ipsum” e post “Hello world!”; aplicar chave de produção do WPML; corrigir variáveis do banner de cookies.",
    "**2. Curto prazo (severidade Média):** completar tradução EN; corrigir o layout/validação do formulário; investigar exceções de JS; tratar a rota /en/ e adicionar hreflang.",
    "**3. UX (revisão da equipe):** unificar ícones do header; corrigir estado do botão do esquema; ajustar proporção do “VEJA MAIS”; melhorar legibilidade nos blocos com blur; remover caixa vazia; adicionar setas ao carrossel; rever timing do pop-up; padronizar proporção/alinhamento das imagens de produto.",
    "**4. Melhoria contínua (severidade Baixa):** padronizar tipografia; ajustar cards de Aplicações; padronizar datas; adicionar alt às imagens.",
    "**5. Reteste pendente:** validar o layout responsivo (mobile), não confirmado nesta avaliação.",
]:
    bullet(it)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Relatório gerado em 17/06/2026 a partir da avaliação do ambiente staging.loccus.com.br. "
              "Documento técnico de uso interno — editável.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = CINZA

doc.save("Relatorio_Tecnico_Avaliacao_staging_loccus.docx")
print("DOCX gerado com sucesso.")
